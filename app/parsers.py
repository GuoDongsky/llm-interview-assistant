from __future__ import annotations

from io import BytesIO
from pathlib import Path

from fastapi import HTTPException


SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".docx"}
SUPPORTED_RESUME_EXTENSIONS = {".docx", ".pdf"}


def _decode_text(content: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise HTTPException(status_code=400, detail="无法识别文件编码，请转换为 UTF-8 后重试。")


def _parse_docx(content: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="服务端缺少 python-docx 依赖，请先安装 requirements.txt。",
        ) from exc

    document = Document(BytesIO(content))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    table_lines: list[str] = []
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))

    return "\n".join(paragraphs + table_lines).strip()


def _parse_pdf(content: bytes) -> str:
    try:
        import fitz
    except ImportError as exc:
        raise HTTPException(
            status_code=500,
            detail="服务端缺少 PyMuPDF 依赖，请先安装 requirements.txt。",
        ) from exc

    try:
        document = fitz.open(stream=content, filetype="pdf")
    except Exception as exc:
        raise HTTPException(status_code=400, detail="PDF 文件无法打开，请检查文件是否损坏。") from exc

    page_text: list[str] = []
    try:
        for page in document:
            text = page.get_text("text").strip()
            if text:
                page_text.append(text)
    finally:
        document.close()

    return "\n\n".join(page_text).strip()


def parse_file_content(filename: str, content: bytes, max_bytes: int) -> tuple[str, str]:
    filename = filename or "uploaded-file"
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="仅支持 txt、md、markdown、docx 文件。")

    if len(content) > max_bytes:
        raise HTTPException(status_code=400, detail="文件过大，当前限制为 10MB。")

    if suffix in {".txt", ".md", ".markdown"}:
        text = _decode_text(content)
    else:
        text = _parse_docx(content)

    text = text.strip()
    if not text:
        raise HTTPException(status_code=400, detail="文件解析后没有可用文本。")

    return filename, text


def parse_resume_content(filename: str, content: bytes, max_bytes: int) -> tuple[str, str]:
    filename = filename or "uploaded-resume"
    suffix = Path(filename).suffix.lower()

    if suffix not in SUPPORTED_RESUME_EXTENSIONS:
        raise HTTPException(status_code=400, detail="简历仅支持 docx、pdf 文件。")

    if len(content) > max_bytes:
        raise HTTPException(status_code=400, detail="文件过大，当前限制为 10MB。")

    if suffix == ".docx":
        text = _parse_docx(content)
    else:
        text = _parse_pdf(content)

    text = text.strip()
    if not text:
        raise HTTPException(
            status_code=400,
            detail="简历解析后没有可用文本。若为扫描版 PDF，请改用可复制文本的 PDF 或 Word 简历。",
        )

    return filename, text
