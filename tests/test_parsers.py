from __future__ import annotations

from io import BytesIO

import fitz
import pytest
from docx import Document
from fastapi import HTTPException

from app.parsers import parse_file_content, parse_resume_content

MAX_BYTES = 10 * 1024 * 1024


def make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for row_index, row in enumerate(table_rows):
            for col_index, cell_text in enumerate(row):
                table.rows[row_index].cells[col_index].text = cell_text
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def make_pdf(text: str) -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    content = document.tobytes()
    document.close()
    return content


def test_parse_txt_utf8():
    filename, text = parse_file_content("记录.txt", "面试官：请自我介绍。".encode("utf-8"), MAX_BYTES)
    assert filename == "记录.txt"
    assert "面试官" in text


def test_parse_txt_gb18030_fallback():
    content = "候选人：我有五年经验。".encode("gb18030")
    _, text = parse_file_content("record.txt", content, MAX_BYTES)
    assert "五年经验" in text


def test_parse_markdown():
    _, text = parse_file_content("notes.md", "# 一面记录\n表现不错".encode("utf-8"), MAX_BYTES)
    assert "一面记录" in text


def test_parse_docx_paragraphs_and_tables():
    content = make_docx(["面试官：谈谈项目经历。"], [["技能", "Python"]])
    _, text = parse_file_content("记录.docx", content, MAX_BYTES)
    assert "项目经历" in text
    assert "技能 | Python" in text


def test_reject_unsupported_extension():
    with pytest.raises(HTTPException) as excinfo:
        parse_file_content("photo.png", b"data", MAX_BYTES)
    assert excinfo.value.status_code == 400


def test_reject_oversize_file():
    with pytest.raises(HTTPException) as excinfo:
        parse_file_content("big.txt", b"x" * 100, max_bytes=10)
    assert excinfo.value.status_code == 400


def test_reject_empty_text():
    with pytest.raises(HTTPException) as excinfo:
        parse_file_content("blank.txt", "   \n  ".encode("utf-8"), MAX_BYTES)
    assert excinfo.value.status_code == 400


def test_parse_resume_pdf():
    content = make_pdf("Python developer with 5 years experience")
    _, text = parse_resume_content("resume.pdf", content, MAX_BYTES)
    assert "Python developer" in text


def test_parse_resume_docx():
    content = make_docx(["姓名：张三", "技能：数据分析"])
    _, text = parse_resume_content("简历.docx", content, MAX_BYTES)
    assert "数据分析" in text


def test_resume_reject_txt():
    with pytest.raises(HTTPException) as excinfo:
        parse_resume_content("resume.txt", b"plain text", MAX_BYTES)
    assert excinfo.value.status_code == 400


def test_resume_reject_corrupt_pdf():
    with pytest.raises(HTTPException) as excinfo:
        parse_resume_content("broken.pdf", b"not a real pdf", MAX_BYTES)
    assert excinfo.value.status_code == 400
